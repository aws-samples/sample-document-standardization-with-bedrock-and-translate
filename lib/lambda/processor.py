import json
import boto3
from botocore.config import Config
import os
import subprocess
import requests
from docx import Document
from claude_prompt import get_claude_prompt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Initialize S3 client
config = Config(connect_timeout=5, read_timeout=60, retries={"total_max_attempts": 20, "mode": "adaptive"})
s3_client = boto3.client('s3', config=config)

# Bedrock config
region = "eu-central-1"
bedrock = boto3.client(
    service_name='bedrock-runtime',
    region_name=region,
    endpoint_url=f'https://bedrock-runtime.{region}.amazonaws.com',
    config=config)

def handler(event, context):
    try:
        # Retrieve bucket name and document key from the event object
        bucket_name = event['documentPath']  
        document_key = event['documentName']  
        reference_key = 'custom-reference.docx'  
        output_bucket = os.environ['OUTPUT_BUCKET']  
        
        # Check if the object key is 'custom-reference.docx'
        if document_key == reference_key:
            return {
                'statusCode': 200,
                'body': json.dumps(f'Uploaded reference template, skipping...')
            }

        # Define local paths for temporary file storage
        local_input_path = '/tmp/' + os.path.basename(document_key)
        local_output_path_html = '/tmp/' + os.path.basename(document_key).replace('.docx', '_corrected.html')
        local_reference_path = '/tmp/' + reference_key
        local_output_path_docx = '/tmp/' + os.path.basename(document_key).replace('.docx', '_corrected.docx')

        # Download the DOCX files from S3 to the local path
        s3_client.download_file(bucket_name, document_key, local_input_path)
        s3_client.download_file(bucket_name, reference_key, local_reference_path)
        
        # Extract title and subtitle before conversion
        title, subtitle = extract_first_two_paragraphs(local_input_path)

        # Convert DOCX to HTML using Pandoc
        subprocess.run([
            'pandoc',
            local_input_path,
            '-o',
            local_output_path_html,
            '-t', 'html',
            '--extract-media=/tmp/media'  # Extract images to a temporary directory
        ], check=True)

        # Read the content of the converted HTML file
        with open(local_output_path_html, 'r') as f:
            HTML_content = f.read()

        # Using Claude v2.1 (update as needed)
        modelID = "anthropic.claude-v2:1"
        
        # Retrieve prompt from claude_prompt.py
        model_prompt = get_claude_prompt(HTML_content)
        
        #Bedock
        llm_model_args = {"prompt": model_prompt, "max_tokens_to_sample": 5000,
                          "stop_sequences": [], "temperature": 0.0, "top_p": 0.9}

        body = json.dumps(llm_model_args)

        response = bedrock.invoke_model(
            body=body,
            modelId=modelID,
            accept='application/json',
            contentType='application/json'
        )

        response = json.loads(response.get("body").read())
        corrected_text = response.get("completion")

        # Write the corrected HTML content to a new file
        with open(local_output_path_html, 'w') as f:
            f.write(corrected_text)
        
        #Uncomment the following line if you also want the intermediary html file to be uploaded to the S3 bucket
        #s3_client.upload_file(local_output_path_html, output_bucket, os.path.basename(local_output_path_html))

        # Convert the corrected HTML content back to a Word document
        subprocess.run([
            'pandoc',
            local_output_path_html,
            '-o',
            local_output_path_docx,
            '-t', 'docx',
            '--resource-path=/tmp/media',  # Ensure Pandoc can find the extracted images
            '--reference-doc', local_reference_path
        ], check=True)
        
        # Load the corrected Word document
        doc = Document(local_output_path_docx)
        
        # Remove the first paragraph if it starts with common Bedrock response phrases
        if doc.paragraphs and (doc.paragraphs[0].text.startswith("Human: ") or doc.paragraphs[0].text.startswith("Here is the text with")):
            p = doc.paragraphs[0]._element
            p.getparent().remove(p)
        
        # Add the title and subtitle back to the beginning of the document
        subtitle_para = doc.paragraphs[0].insert_paragraph_before(subtitle, style='Subtitle')
        title_para = doc.paragraphs[0].insert_paragraph_before(title, style='Title')

        # Center all images in the document
        center_images(doc)

        # Save the modified document
        doc.save(local_output_path_docx)

        # Upload the corrected Word document to the specified output S3 bucket
        with open(local_output_path_docx, 'rb') as f:
            s3_client.upload_fileobj(f, output_bucket, local_output_path_docx.split('/')[-1])

        # Cleanup local files
        os.remove(local_input_path)
        os.remove(local_output_path_html)
        os.remove(local_output_path_docx)

        return {
            'statusCode': 200,
            'body': json.dumps(f'Grammar check was successful for {document_key}')
        }
    except Exception as e:
        print(f'Error: {str(e)}')
        return {
            'statusCode': 500,
            'body': json.dumps(f'Could not process {document_key}')
        }

def extract_first_two_paragraphs(document_path):
    doc = Document(document_path)
    paragraphs = [p.text for p in doc.paragraphs[:2]]
    return paragraphs[0], paragraphs[1] if len(paragraphs) > 1 else ""

def center_images(doc):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if 'graphic' in run.element.xml:
                # Center the image by setting the alignment of the parent paragraph
                align_paragraph_center(paragraph)

def align_paragraph_center(paragraph):
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
