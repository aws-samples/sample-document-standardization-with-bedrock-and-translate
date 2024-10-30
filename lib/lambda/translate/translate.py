# Copyright Amazon.com, Inc. or its affiliates. All Rights Reserved.
# SPDX-License-Identifier: MIT-0

import boto3
import os
import json
import docx
import tempfile

s3 = boto3.client('s3')
translate = boto3.client('translate')

## Update the below variables when changing the languages used
LANGUAGE_FOLDERS = ['english', 'spanish', 'french']  
LANGUAGE_CODES = {
    'english': 'en',
    'spanish': 'es',
    'french': 'fr'
}

def create_folder_if_not_exists(bucket_name, folder_name):
    try:
        s3.head_object(Bucket=bucket_name, Key=f'{folder_name}/')
    except s3.exceptions.ClientError as e:
        if e.response['Error']['Code'] == '404':
            s3.put_object(Bucket=bucket_name, Key=f'{folder_name}/')

def determine_language(key):
    for folder in LANGUAGE_FOLDERS:
        if key.startswith(f'{folder}/'):
            return folder
    return None

def translate_text(text, source_language, target_language):
    response = translate.translate_text(
        Text=text,
        SourceLanguageCode=source_language,
        TargetLanguageCode=target_language
    )
    return response['TranslatedText']


def handler(event, context):
    try: 
        bucket_name = event['documentPath']
        document_key = event['documentName']
        original_filename = os.path.basename(document_key)

        all_files = []

        # Check and create language folders if necessary
        for folder in LANGUAGE_FOLDERS:
            create_folder_if_not_exists(bucket_name, folder)
        print("Checked and created language folders if they didn't exist.")

        language_code = determine_language(document_key)
        if language_code is None:
            print(f"Could not determine the source language from the key: {document_key}")
            return
        
        language_code = document_key.split('/')[0]  

        path_dict = {
                'name': original_filename,  
                'path': document_key,
                'language_code': language_code
            }
        all_files.append(path_dict)

        # Translate text and upload to corresponding folders
        for target_folder in LANGUAGE_FOLDERS:
            if target_folder == language_code:
                continue  

            target_language_code = LANGUAGE_CODES[target_folder]
            source_language_code = LANGUAGE_CODES[language_code]            
            
            # Create a temporary file
            with tempfile.NamedTemporaryFile(delete=False,  suffix='.docx') as temp_file:
                download_path = temp_file.name
            
            s3.download_file(bucket_name, document_key, download_path)

            # Load the DOCX file
            doc = docx.Document(download_path)

            # Translate text in paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    translated_text = translate_text(paragraph.text, source_language_code, target_language_code)
                    paragraph.text = translated_text

            # Translate text in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text:
                                translated_text = translate_text(paragraph.text, source_language_code, target_language_code)
                                paragraph.text = translated_text

            # Save the translated document
            with tempfile.NamedTemporaryFile(delete=False,  suffix='.docx') as temp_translated:
                translated_path = temp_translated.name

            doc.save(translated_path)


            # Upload the translated document to the input bucket under the translaed path
            original_filename_without_doctype = original_filename.split('.')[0]
            target_key = f'{target_folder}/{original_filename_without_doctype}_{language_code}_to_{target_folder}_translated.docx' # matches the exempted prefix in the s3EventRule
            input_bucket = os.environ['INPUT_BUCKET']
            s3.upload_file(
                translated_path, 
                input_bucket, 
                target_key
            )
            
            path_dict = {
                'name': f'{original_filename_without_doctype}_{language_code}_to_{target_folder}_translated.docx',  
                'path': target_key,
                'language_code': target_language_code
            }
            
            all_files.append(path_dict)
            os.unlink(translated_path)
            os.unlink(download_path)



        print(f"Successfully processed and translated {target_key}")



        return {
            'statusCode': 200,
            'body': json.dumps({
                'message': 'Translation successful',
                'filePaths': all_files,
                'inputBucket': input_bucket
            })
        }
    except Exception as e:
        return {
            'statusCode': 500,
            "message": f'translate lambda failed due to the following error: {str(e)}'
        }
    



