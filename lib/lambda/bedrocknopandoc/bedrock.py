# Copyright Amazon.com, Inc. or its affiliates. All Rights Reserved.
# SPDX-License-Identifier: MIT-0

import json
import boto3
from botocore.config import Config
import os
import subprocess
from docx import Document
from prompt import get_claude_prompt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import zipfile
import mammoth
from bs4 import BeautifulSoup
from io import BytesIO
import base64
from docx.shared import Pt, RGBColor
import zipfile



# Initialize S3 client
config = Config(connect_timeout=5, read_timeout=60, retries={"total_max_attempts": 20, "mode": "adaptive"})
s3_client = boto3.client('s3', config=config)

# Bedrock config
region = "us-east-1"
bedrock = boto3.client(
    service_name='bedrock-runtime',
    region_name=region,
    endpoint_url=f'https://bedrock-runtime.{region}.amazonaws.com',
    config=config)

def handler(event, context):
    try:
        # Retrieve bucket name and document key from the event object
        bucket_name = os.environ['INPUT_BUCKET']  
       # document_key = event['path']  
        reference_key = 'word_template.docx'  
        document_key = 'english/tone_test.docx'
        output_bucket = os.environ['OUTPUT_BUCKET']  

        # Define local paths for temporary file storage
        local_input_path = '/tmp/' + os.path.basename(document_key)
        local_reference_path = '/tmp/' + os.path.basename(reference_key)
        local_output_path_docx = '/tmp/' + os.path.basename(document_key).replace('.docx', '_corrected.docx')
        tmp_dir = "/tmp/output_images"

        # Download the DOCX file from S3 to the local path
        s3_client.download_file(bucket_name, document_key, local_input_path)
        print('file downloaded')

        # Download the reference file from S3 to the local path
        s3_client.download_file(bucket_name, reference_key, local_reference_path)
        print('reference file downloaded')

        # Extract images and replace them with placeholders
        images_info = extract_images_and_replace_with_placeholders(local_input_path, tmp_dir)
        print('images extracted and replaced with placeholders')
        print(images_info)

        #Convert DOCX to HTML using Mammoth
        html_content = docx_to_html(local_input_path)
        print('html content generated')
        print(html_content)
        
        # Retrieve prompt from claude_prompt.py
        model_prompt = get_claude_prompt(html_content)

        native_request = {
            "anthropic_version": "bedrock-2023-05-31",
            "max_tokens": 5000,
            "temperature": 0.0,
            "messages": [
                {
                    "role": "user",
                    "content": [{"type": "text", "text": model_prompt}],
                }
            ],
        }

        body = json.dumps(native_request)

        # Using Claude 3 Sonnet (update as needed)
        modelID = "anthropic.claude-3-sonnet-20240229-v1:0"
        
        response = bedrock.invoke_model(
            body=body,
            modelId=modelID,
        )

        response = json.loads(response.get("body").read())
        corrected_text = response["content"][0]["text"]
        print('text corrected')
        print(corrected_text)

        #html_to_docx(corrected_text, local_output_path_docx)

        #loading template and adding stuff to it

        load_template_and_add_html_content(local_reference_path, local_output_path_docx, corrected_text)

        reinsert_images(local_output_path_docx, images_info)


        print(images_info)

        print('docx generated')
        
        # Load the corrected Word document
        doc = Document(local_output_path_docx)

        # Center all images in the document
        center_images(doc)

        # Save the modified document
        doc.save(local_output_path_docx)

        if document_key.endswith("_translated.docx"):
            final_doc_name = document_key.replace('_translated.docx', '_corrected.docx')
        else:
            final_doc_name = document_key.replace('.docx', '_corrected.docx')

        # Upload the corrected Word document to the specified output S3 bucket
        with open(local_output_path_docx, 'rb') as f:
            s3_client.upload_fileobj(f, output_bucket, final_doc_name)

        # Cleanup local files
        os.remove(local_input_path)
        os.remove(local_output_path_docx)

        return {
            'statusCode': 200,
            'body': final_doc_name
        }
    except Exception as e:
        print(f'Error: {str(e)}')
        return {
            'statusCode': 500,
            'body': json.dumps(f'Could not process {document_key} due to the following error: {str(e)}')
        }
    


##USING##
def center_images(doc):
    """Center images in doc."""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if 'graphic' in run.element.xml:
                align_paragraph_center(paragraph)
##USING##

def align_paragraph_center(paragraph):
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)

def docx_to_html(docx_path):
    """Convert DOCX to HTML using Mammoth."""
    with open(docx_path, "rb") as docx_file:      
        html_content = mammoth.convert_to_html(docx_file).value
        return html_content


def insert_image(element, doc):
    """Handle the insertion of images."""
    base64_data = element['src']
    
    # Ensure the Base64 data has the correct format (starts with "data:image")
    if base64_data.startswith("data:image"):
        # Strip the metadata and only get the image data
        base64_data = base64_data.split(",")[1]  # Get only the Base64 data
        
        try:
            # Decode the Base64 data
            image_data = base64.b64decode(base64_data)
            
            # Convert the decoded image data into a BytesIO object for python-docx
            image_stream = BytesIO(image_data)
            
            # Add the image to the document (optionally adjust the size)
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(image_stream) 
            
        except Exception as e:
            print(f"Error decoding or inserting image: {e}")

def _style_heading(text, doc, level, font_size, color=None, italic=False, bold=False):
    """Helper function to apply styling to headings (used in mapping)."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(font_size)
    
    if color:
        run.font.color.rgb = RGBColor(*color)
    
    run.italic = italic  
    run.bold = bold

def _style_text(element, run):
    """Apply styles like bold and italic to a run."""
    # Apply bold and italic based on the element's tag
    if element.name in ['strong', 'b']:
        run.bold = True
    if element.name in ['em', 'i']:
        run.italic = True
    return run

# def _add_list_item(element, doc, list_type="unordered"):
#     """Add list items (ul/ol) to the document."""
#     # Adjust the indentation and style for lists
#     if list_type == "unordered":
#         paragraph = doc.add_paragraph(element.text, style='ListBullet')
#     elif list_type == "ordered":
#         paragraph = doc.add_paragraph(element.text, style='ListNumber')
#     return paragraph

# def _add_list_item(element, doc, list_type="unordered"):
#     """Add list items (ul/ol) to the document, handling nested lists."""
#     if list_type == "unordered":
#         paragraph = doc.add_paragraph(element.text, style='ListBullet')
#     elif list_type == "ordered":
#         paragraph = doc.add_paragraph(element.text, style='ListNumber')
    
#     # Check for nested lists (children of the list item)
#     for child in element.children:
#         if child.name == "ul":
#             for sub_item in child.find_all("li"):
#                 _add_list_item(sub_item, doc, list_type="unordered")
#         elif child.name == "ol":
#             for sub_item in child.find_all("li"):
#                 _add_list_item(sub_item, doc, list_type="ordered")

#     return paragraph

def html_to_docx(corrected_html, docx_file_path):
    """Convert corrected HTML back to .docx and reinsert Base64-encoded images."""
    doc = Document()
    soup = BeautifulSoup(corrected_html, "html.parser")

    # Inline styling for h1 and other elements in the mapping
    html_to_docx_mapping = {
        'p': lambda element, doc: doc.add_paragraph(element.text),
        'h1': lambda element, doc: _style_heading(element.text, doc, level=1, font_size=24, color=(79, 129, 189), italic=True),  # Blue, italic, larger font for h1
        'h2': lambda element, doc: _style_heading(element.text, doc, level=2, font_size=18), #h2
        'h3': lambda element, doc: _style_heading(element.text, doc, level=3, font_size=16), #h3
        'h4': lambda element, doc: _style_heading(element.text, doc, level=4, font_size=14, italic=True), #h4 
        'ul': lambda element, doc: _add_list_item(element, doc, list_type="unordered"), #unordered bullet points
        'ol': lambda element, doc: _add_list_item(element, doc, list_type="ordered"), #ordered bullet points
        'strong': lambda element, doc: _style_text(element, doc.add_paragraph().add_run()), #bold
        'b': lambda element, doc: _style_text(element, doc.add_paragraph().add_run()), #bold
        'em': lambda element, doc: _style_text(element, doc.add_paragraph().add_run()), #italic
        'i': lambda element, doc: _style_text(element, doc.add_paragraph().add_run()) #italic
    }

    # Iterate over all elements in the soup
    for element in soup:
        # Check if the element is in the mapping dictionary
        if element.name in html_to_docx_mapping:
            # Call the corresponding function (with styling, image handling, etc.)
            html_to_docx_mapping[element.name](element, doc)

    # Save the new .docx file
    doc.save(docx_file_path)


def extract_images_and_replace_with_placeholders(docx_file_path, tmp_dir):
    """Extract images from Word document, save them to a temporary directory, and replace with placeholders."""
    doc = Document(docx_file_path)
    images_info = []
    image_counter = 1

    # Ensure tmp directory exists
    os.makedirs(tmp_dir, exist_ok=True)

    # Open the DOCX file as a ZIP archive to extract images
    with zipfile.ZipFile(docx_file_path, 'r') as docx_zip:
        # Loop through all paragraphs and runs to find images (graphics)
        for i, paragraph in enumerate(doc.paragraphs):
            for run in paragraph.runs:
                if 'graphic' in run.element.xml:
                    # Create a placeholder and file path for the image
                    placeholder = f'[IMAGE_{image_counter}]'
                    image_filename = f'image_{image_counter}.png'
                    image_path = os.path.join(tmp_dir, image_filename)

                    # Extract image data from the DOCX ZIP archive
                    for rel in doc.part.rels:
                        if "image" in doc.part.rels[rel].target_ref:
                            image = doc.part.rels[rel].target_part
                            image_bytes = image._blob

                            # Save the image to the temporary directory
                            with open(image_path, 'wb') as img_file:
                                img_file.write(image_bytes)

                            # Store the image information
                            images_info.append({
                                "placeholder": placeholder,
                                "image_path": image_path,
                                "paragraph_index": i
                            })

                            # Replace the image with a placeholder in the paragraph
                            paragraph.clear()
                            paragraph.add_run(placeholder)

                            image_counter += 1
                            break  # Stop after processing one image per run

    # Save the modified DOCX with placeholders
    doc.save(docx_file_path)
    return images_info


def reinsert_images(docx_file_path, images_info):
    """Reinsert images in place of placeholders in the Word document."""
    doc = Document(docx_file_path)
    for image_info in images_info:
        placeholder = image_info["placeholder"]
        print(f'placeholder : {placeholder}')
        image_path = image_info["image_path"]
        print(f'image_path : {image_path}')
        paragraph_index = image_info["paragraph_index"]

        # Loop through all paragraphs in the document
        for paragraph in doc.paragraphs:
            # Check if the placeholder exists in the paragraph
            if placeholder in paragraph.text:
                # Replace the placeholder with the image
                paragraph.clear()  # Clear the placeholder text
                run = paragraph.add_run()  # Create a new run to insert the image
                run.add_picture(image_path)  # Insert the image 

    # Save the final DOCX file with images reinserted
    doc.save(docx_file_path)

def _add_list(element, doc, level, list_type):
    """
    Recursively process ordered/unordered lists and apply styles based on nesting level.
    """
    # Select the appropriate list style based on the nesting level and list type
    if list_type == 'unordered':
        list_style = 'ListBullet' if level == 0 else f'ListBullet{level + 1}'  # ListBullet, ListBullet2, etc.
    else:
        list_style = 'ListNumber' if level == 0 else f'ListNumber{level + 1}'  # ListNumber, ListNumber2, etc.

    # Loop through the immediate children <li> elements of the current list
    for li in element.find_all('li', recursive=False):
        # Get the text for the current list item
        list_text = ""
        for child in li.children:
            if child.name is None:
                list_text += child.strip() + " "

        # Add the list item to the document with the correct style
        paragraph = doc.add_paragraph(list_text.strip(), style=list_style)

        # Process any nested <ul> or <ol> inside this <li> recursively
        nested_list = li.find(['ul', 'ol'])
        if nested_list:
            # Determine whether the nested list is ordered or unordered
            nested_list_type = 'unordered' if nested_list.name == 'ul' else 'ordered'
            _add_list(nested_list, doc, level + 1, nested_list_type)


def clear_document_body(doc):
    """Remove all content from the document body while preserving headers, footers, and styles."""
    # Remove all paragraphs and content from the document body
    for paragraph in doc.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)

    # Remove all tables, if any exist
    for table in doc.tables:
        tbl = table._element
        tbl.getparent().remove(tbl)

    # The headers and footers remain intact
    return doc

def load_template_and_add_html_content(template_path, output_path, html_content):
    """Load a pre-styled template DOCX, add content from HTML, and save it to S3."""
    # Load the pre-styled template DOCX from S3
    doc = Document(template_path)
    print(f'Loaded template')

    # Clear the body of the template
    clear_document_body(doc)
    print(f'Cleared body')

    # Parse the HTML content using BeautifulSoup
    soup = BeautifulSoup(html_content, 'html.parser')

    # HTML-to-DOCX mapping for common elements and styles in the template
    html_to_docx_mapping = {
        'p': lambda element, doc: doc.add_paragraph(element.text, style='Normal'),
        'h1': lambda element, doc: doc.add_paragraph(element.text, style='Heading 1'),
        'h2': lambda element, doc: doc.add_paragraph(element.text, style='Heading 2'),
        'h3': lambda element, doc: doc.add_paragraph(element.text, style='Heading 3'),
        'h4': lambda element, doc: doc.add_paragraph(element.text, style='Heading 4'),
        'ul': lambda element, doc: _add_list(element, doc, list_type='unordered', level=0),
        'ol': lambda element, doc: _add_list(element, doc, list_type='ordered', level=0),
        'strong': lambda element, doc: _style_text(element, doc.add_paragraph().add_run()), #bold
        'b': lambda element, doc: _style_text(element, doc.add_paragraph().add_run()), #bold
        'em': lambda element, doc: _style_text(element, doc.add_paragraph().add_run()), #italic
        'i': lambda element, doc: _style_text(element, doc.add_paragraph().add_run()) #italic
    }

    # Iterate over all HTML elements and apply the mapping
    for element in soup:
        if element.name in html_to_docx_mapping:
            html_to_docx_mapping[element.name](element, doc)

    # Save the modified document to the output bucket
    doc.save(output_path)

