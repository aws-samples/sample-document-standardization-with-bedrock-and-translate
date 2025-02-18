# Copyright Amazon.com, Inc. or its affiliates. All Rights Reserved.
# SPDX-License-Identifier: MIT-0

import json
import boto3
from botocore.config import Config
import os
from docx import Document
from claude_prompt import get_claude_prompt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.oxml.shared
from docx.opc.constants import RELATIONSHIP_TYPE
import mammoth
from bs4 import BeautifulSoup
from docx.shared import Pt
import zipfile
import tempfile
import shutil


# Initialize S3 client
config = Config(connect_timeout=5, read_timeout=60, retries={"total_max_attempts": 20, "mode": "adaptive"})
s3_client = boto3.client('s3', config=config)

# Bedrock config
region = "us-east-1"
bedrock = boto3.client(
    service_name='bedrock-runtime',
    region_name=region,
    endpoint_url=f'https://bedrock-runtime.{region}.amazonaws.com',
    config=config)

def handler(event, context):
    try:
        # Retrieve bucket name and document key from the event object
        bucket_name = os.environ['INPUT_BUCKET']  
        document_key = event['path']  
        reference_key = 'word_template.docx'  
        output_bucket = os.environ['OUTPUT_BUCKET']
  
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_input:
            local_input_path = temp_input.name
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_reference:
            local_reference_path = temp_reference.name
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_output:
            local_output_path_docx = temp_output.name
        
        # Create a temporary directory
        tmp_dir = tempfile.mkdtemp(prefix='output_images_')
        

        # Download the DOCX file from S3 to the local path
        s3_client.download_file(bucket_name, document_key, local_input_path)

        # Download the reference template from S3 to the local path
        s3_client.download_file(bucket_name, reference_key, local_reference_path)

        # Extract images and replace them with placeholders
        images_info = extract_images_and_replace_with_placeholders(local_input_path, tmp_dir)
        
        #Convert DOCX to HTML using Mammoth
        html_content = docx_to_html(local_input_path)
        
        # Retrieve prompt from claude_prompt.py
        model_prompt = get_claude_prompt(html_content)

        # Send HTML to model for processing
        corrected_text = invoke_bedrock_model(model_prompt)

        # loading template and transforming HTML back to DOCX
        load_template_and_add_html_content(local_reference_path, local_output_path_docx, corrected_text)

        # reinstering images that were removed
        reinsert_images(local_output_path_docx, images_info)
        
        # Load the corrected Word document
        doc = Document(local_output_path_docx)

        # Center all images in the document
        center_images(doc)

        # Save the modified document
        doc.save(local_output_path_docx)

        if document_key.endswith("_translated.docx"):
            final_doc_name = document_key.replace('_translated.docx', '_corrected.docx')
        else:
            final_doc_name = document_key.replace('.docx', '_corrected.docx')
        
        
        # Upload the corrected Word document to the specified output S3 bucket
        with open(local_output_path_docx, 'rb') as f:
            s3_client.upload_fileobj(f, output_bucket, final_doc_name)


        # Clean up temporary files
        os.unlink(local_input_path)
        os.unlink(local_reference_path)
        os.unlink(local_output_path_docx)

        # Clean up temporary directory
        shutil.rmtree(tmp_dir)

        return {
            'statusCode': 200,
            'body': final_doc_name
        }
    except Exception as e:
        print(f'Error: {str(e)}')
        return {
            'statusCode': 500,
            'body': json.dumps(f'Could not process {document_key} due to the following error: {str(e)}')
        }
    

## Functions used above##

def invoke_bedrock_model(model_prompt):
    """Invoke Bedrock model."""
    
    native_request = {
            "anthropic_version": "bedrock-2023-05-31",
            "max_tokens": 5000,
            "temperature": 0.0,
            "messages": [
                {
                    "role": "user",
                    "content": [{"type": "text", "text": model_prompt}],
                }
            ],
        }

    body = json.dumps(native_request)

    # Using Claude 3 Sonnet (update as needed)
    modelID = "anthropic.claude-3-sonnet-20240229-v1:0"
    
    response = bedrock.invoke_model(
        body=body,
        modelId=modelID,
    )

    response = json.loads(response.get("body").read())
    corrected_text = response["content"][0]["text"]
    return corrected_text

def center_images(doc):
    """Center images in doc."""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if 'graphic' in run.element.xml:
                align_paragraph_center(paragraph)

def align_paragraph_center(paragraph):
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)


def docx_to_html(docx_path):
    """Convert DOCX to HTML using Mammoth."""
    with open(docx_path, "rb") as docx_file:      
        html_content = mammoth.convert_to_html(docx_file).value
        print(html_content)
        return html_content


def _style_text(element, run):
    """Apply styles like bold and italic to a run."""
    # Apply bold and italic based on the element's tag
    if element.name in ['strong', 'b']:
        run.bold = True
    if element.name in ['em', 'i']:
        run.italic = True
    return run

def extract_images_and_replace_with_placeholders(docx_file_path, tmp_dir):
    """Extract images from Word document, save them to a temporary directory, and replace with placeholders."""
    doc = Document(docx_file_path)
    images_info = []
    image_counter = 1
    
    # Ensure tmp directory exists
    os.makedirs(tmp_dir, exist_ok=True)
    
    # Keep track of processed relationships to avoid duplicates
    processed_rels = set()

    with zipfile.ZipFile(docx_file_path, 'r') as docx_zip:
        for i, paragraph in enumerate(doc.paragraphs):
            for run in paragraph.runs:
                if 'graphic' in run.element.xml:
                    # Find the relationship ID for this specific graphic
                    graphic_element = run.element.find('.//a:graphic', 
                        namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
                    if graphic_element is None:
                        continue
                        
                    # Get the relationship ID for this specific image
                    blip_element = graphic_element.find('.//a:blip', 
                        namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
                    if blip_element is None:
                        continue
                        
                    # Get the relationship ID (rId) for this image
                    rel_id = blip_element.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    
                    # Skip if we've already processed this relationship
                    if rel_id in processed_rels:
                        continue
                    processed_rels.add(rel_id)

                    # Create placeholder and file path
                    placeholder = f'[IMAGE_{image_counter}]'
                    image_filename = f'image_{image_counter}.png'
                    image_path = os.path.join(tmp_dir, image_filename)

                    # Get the specific image using the relationship ID
                    if rel_id in doc.part.rels:
                        image = doc.part.rels[rel_id].target_part
                        image_bytes = image._blob

                        # Save the image to the temporary directory
                        with open(image_path, 'wb') as img_file:
                            img_file.write(image_bytes)

                        # Store the image information
                        images_info.append({
                            "placeholder": placeholder,
                            "image_path": image_path,
                            "paragraph_index": i,
                            "rel_id": rel_id  # Store the rel_id for debugging
                        })

                        # Replace with placeholder
                        run.clear()
                        run.text = placeholder
                        
                        image_counter += 1

    # Save the modified DOCX with placeholders
    doc.save(docx_file_path)
    
    # Print debug information
    print(f"Processed {len(processed_rels)} unique images")
    # for info in images_info:
    #     print(f"Image: {info['image_path']}, RelID: {info['rel_id']}")
        
    return images_info

def reinsert_images(docx_file_path, images_info):
    doc = Document(docx_file_path)
    
    # Create a mapping of placeholder to image info
    placeholder_map = {info["placeholder"]: info for info in images_info}
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # Check if this run contains a placeholder
            for placeholder, info in list(placeholder_map.items()):  
                if placeholder in run.text:
                    run.clear()
                    run.add_picture(info["image_path"])
                    # Remove this placeholder from the map to avoid reusing
                    del placeholder_map[placeholder]
                    break

    doc.save(docx_file_path)

def _add_list(element, doc, level, list_type):
    """
    Recursively process ordered/unordered lists and apply styles based on nesting level.
    """
    # Select the appropriate list style based on the nesting level and list type
    if list_type == 'unordered':
        list_style = 'ListBullet' if level == 0 else f'ListBullet{level + 1}'  # ListBullet, ListBullet2, etc.
    else:
        list_style = 'ListNumber' if level == 0 else f'ListNumber{level + 1}'  # ListNumber, ListNumber2, etc.

    # Loop through the immediate children <li> elements of the current list
    for li in element.find_all('li', recursive=False):
        # Get the text for the current list item
        list_text = ""
        for child in li.children:
            if child.name is None:
                list_text += child.strip() + " "

        # Add the list item to the document with the correct style
        paragraph = doc.add_paragraph(list_text.strip(), style=list_style)

        # Process any nested <ul> or <ol> inside this <li> recursively
        nested_list = li.find(['ul', 'ol'])
        if nested_list:
            # Determine whether the nested list is ordered or unordered
            nested_list_type = 'unordered' if nested_list.name == 'ul' else 'ordered'
            _add_list(nested_list, doc, level + 1, nested_list_type)


def clear_document_body(doc):
    """Remove all content from the document body of the template while preserving headers, footers, and styles."""
    # Remove all paragraphs and content from the document body
    for paragraph in doc.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)

    # Remove all tables, if any exist
    for table in doc.tables:
        tbl = table._element
        tbl.getparent().remove(tbl)

    return doc

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = docx.oxml.shared.OxmlElement('w:r')
    
    # Create a new text object (a wrapper over a 'w:t' element)
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    
    # Add color
    c = docx.oxml.shared.OxmlElement('w:color')
    c.set(docx.oxml.shared.qn('w:val'), '0000FF')
    rPr.append(c)
    
    # Add underline
    u = docx.oxml.shared.OxmlElement('w:u')
    u.set(docx.oxml.shared.qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_text = docx.oxml.shared.OxmlElement('w:t')
    new_text.text = text
    new_run.append(new_text)
    hyperlink.append(new_run)
    
    # Add the hyperlink to the paragraph
    paragraph._p.append(hyperlink)

def load_template_and_add_html_content(template_path, output_path, html_content):
    """Load a pre-styled template DOCX, add content from HTML, and save it to S3."""
    
    doc = Document(template_path)
    clear_document_body(doc)
    soup = BeautifulSoup(html_content, 'html.parser')

    def process_paragraph_content(paragraph, element):
        """Process the content of a paragraph, handling nested formatting."""
        current_text = ''
        
        for content in element.contents:
            if isinstance(content, str):
                if current_text:
                    paragraph.add_run(current_text)
                current_text = content
            else:
                if current_text:
                    paragraph.add_run(current_text)
                    current_text = ''
                
                if content.name in ['strong', 'b']:
                    run = paragraph.add_run(content.get_text())
                    run.bold = True
                elif content.name in ['em', 'i']:
                    run = paragraph.add_run(content.get_text())
                    run.italic = True
                elif content.name == 'a':
                    # Add the hyperlink
                    href = content.get('href', '')
                    if href:
                        add_hyperlink(paragraph, content.get_text(), href)
                else:
                    paragraph.add_run(content.get_text())
        
        if current_text:
            paragraph.add_run(current_text)

    # HTML-to-DOCX mapping for common elements and styles in the template
    html_to_docx_mapping = {
        'p': lambda element, doc: process_paragraph_content(doc.add_paragraph(), element),
        'ul': lambda element, doc: _add_list(element, doc, list_type='unordered', level=0),
        'ol': lambda element, doc: _add_list(element, doc, list_type='ordered', level=0),
        'strong': lambda element, doc: _style_text(element, doc.add_paragraph().add_run()),
        'b': lambda element, doc: _style_text(element, doc.add_paragraph().add_run()),
        'em': lambda element, doc: _style_text(element, doc.add_paragraph().add_run()),
        'i': lambda element, doc: _style_text(element, doc.add_paragraph().add_run())
    }

    # Add all possible header levels
    html_to_docx_mapping.update({
        f'h{i}': lambda element, doc, i=i: process_paragraph_content(
            doc.add_paragraph(style=f'Heading {i}'), 
            element
        ) for i in range(1, 10)
    })

    # Iterate over all HTML elements and apply the mapping
    for element in soup:
        if element.name in html_to_docx_mapping:
            html_to_docx_mapping[element.name](element, doc)

    doc.save(output_path)
